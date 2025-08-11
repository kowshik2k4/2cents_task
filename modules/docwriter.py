from docx import Document
import tempfile

def create_reviewed_doc(filename, text, issues):
    # Create a new Document with the text + reviewer comments
    doc = Document()
    doc.add_heading(f"Reviewed: {filename}", level=1)
    doc.add_paragraph(text)
    
    if issues:
        doc.add_heading("Reviewer Comments", level=2)
        for issue in issues:
            comment_text = f"[{issue['severity']}] {issue['issue']} — Suggestion: {issue['suggestion']}"
            doc.add_paragraph(comment_text)
    
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return tmp.name
